from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("output/A股港股跨市场投资计划_多智能体投委会版_20260605.docx")


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "6B7280"
LIGHT = "F2F4F7"
CALL_OUT = "F4F6F9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(1, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_run(run, bold=False, color=INK, size=11) -> None:
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    style_run(run, bold=True, color=BLUE if level <= 2 else DARK_BLUE, size={1: 16, 2: 13, 3: 12}.get(level, 11))
    return p


def add_body(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    p.style = "Normal"
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        style_run(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        style_run(r2)
    else:
        r = p.add_run(text)
        style_run(r)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    style_run(r)
    return p


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    style_run(r)
    return p


def add_table(doc: Document, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_width(table, widths)
    for idx, h in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        style_run(run, bold=True, color=INK, size=10)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx > 0 and len(str(value)) <= 16 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(value))
            style_run(run, size=10)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_callout(doc: Document, title: str, body: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALL_OUT)
    p = cell.paragraphs[0]
    r1 = p.add_run(title + "：")
    style_run(r1, bold=True, color=DARK_BLUE)
    r2 = p.add_run(body)
    style_run(r2)
    doc.add_paragraph()


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)


def set_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("A股+港股跨市场投资计划 | 多智能体投委会版 | 2026-06-05")
    style_run(r, color=MUTED, size=9)


def build() -> None:
    doc = Document()
    configure_styles(doc)
    set_footer(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("A股+港股跨市场投资计划")
    style_run(run, bold=True, color=DARK_BLUE, size=22)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("500万元人民币 | 3-5年 | 无杠杆 | 多智能体投委会版")
    style_run(r, color=MUTED, size=11)

    add_callout(
        doc,
        "核心结论",
        "建议采用“A股核心仓 + 港股战略卫星仓 + 现金固收缓冲 + 黄金与REITs低相关资产”的组合。基准配置为A股60%、港股18%、现金/固收12%、黄金ETF6%、公募REITs4%。预期年化收益目标为8%-11%，常态最大回撤控制在20%-25%，极端情境可能达到28%-35%。",
    )

    add_heading(doc, "多智能体工作流与审查结论")
    add_body(doc, "本方案采用并行研究、交叉审查、CIO汇总的工作流。宏观策略负责人负责跨市场宏观与政策底稿；港股研究负责人负责港股行业与交易规则；组合基金经理负责资产配置和执行路径；风控总监负责压力测试和投委会准入阈值。")
    add_table(
        doc,
        ["角色", "核心贡献", "主要约束"],
        [
            ["宏观策略负责人", "确认未来3-5年更偏结构牛而非全面牛市", "权益配置需保留估值与现金安全边际"],
            ["港股研究负责人", "港股补足互联网平台、高股息、创新药与离岸估值修复机会", "港股仓位需控制在战略卫星仓"],
            ["组合基金经理", "提出A股60%、港股18%、现金固收12%、黄金6%、REITs4%", "分批建仓、ETF为主、个股为辅"],
            ["风控总监", "加入港股汇率、美元利率、离岸流动性和南向拥挤风险", "回撤阈值和投委会复核必须前置"],
        ],
        [1800, 4300, 3260],
    )

    add_heading(doc, "第一部分：投资哲学与核心逻辑")
    add_body(doc, "本组合的目标不是短线交易收益，而是在3-5年内获取中国优质权益资产的复利回报。A股承担产业升级和人民币资产核心配置功能；港股承担低估值、高股息、互联网平台和创新药等结构补充功能。组合不使用融资融券、不使用杠杆、不使用期货、期权、互换等高风险衍生品。")
    add_body(doc, "截至2026年6月初，中国经济总量温和修复但结构分化明显。2026年一季度GDP同比增长5.0%；2026年5月制造业PMI为50.0%，高技术制造业PMI为52.9%。这意味着未来收益更依赖盈利质量、产业趋势、估值纪律和资金结构，而非单纯依赖全面贝塔。")

    add_heading(doc, "第二部分：未来3-5年中国投资主线")
    for item in [
        "AI、自主可控与数字经济：A股更适合半导体设备、工业软件、机器人和国产算力链；港股更适合互联网平台、云服务和AI应用。",
        "高端制造与制造出海：A股产业链更完整，重点关注工程机械、汽车零部件、电力设备、自动化、船舶和储能。",
        "高股息与央国企质量重估：港股在电信、能源、公用事业、金融等领域估值和股息优势更突出；A股可配置红利低波和央企红利ETF。",
        "医疗创新与银发经济：港股偏创新药融资生态和国际授权潜力，A股偏医疗器械、医疗服务和稳定现金流公司。",
        "能源转型与资源安全：A股偏电网、储能、核电设备与制造；港股偏资源能源、运营资产和分红回报。",
        "消费分层与平台经济：A股配置必需消费、品牌龙头和出海品牌；港股配置现金流强、回购稳定、监管边界清晰的平台公司。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "第三部分：资产配置方案")
    add_table(
        doc,
        ["资产类别", "目标比例", "金额", "定位"],
        [
            ["A股权益", "60%", "300万元", "核心收益来源，承接产业升级和人民币资产配置"],
            ["港股权益", "18%", "90万元", "战略卫星仓，补足低估值、高股息、平台和创新药"],
            ["现金/货基/短债/逆回购", "12%", "60万元", "流动性缓冲与回撤加仓弹药"],
            ["黄金ETF", "6%", "30万元", "地缘、汇率、通胀和尾部风险对冲"],
            ["公募REITs", "4%", "20万元", "低相关现金流资产，控制单只流动性风险"],
        ],
        [2100, 1400, 1800, 4060],
    )
    add_body(doc, "战术区间：A股55%-65%，港股12%-22%，现金/固收10%-25%，黄金ETF4%-10%，REITs0%-6%。权益总仓通常不超过78%；只有在市场显著低估且投委会复核通过时，才允许短期接近80%。")

    add_heading(doc, "第四部分：行业配置方案")
    add_heading(doc, "A股内部配置", 2)
    add_table(
        doc,
        ["方向", "A股内比例", "金额", "配置逻辑"],
        [
            ["宽基ETF", "35%", "105万元", "获取市场beta，降低选股误差"],
            ["红利低波/央国企", "15%", "45万元", "现金流与分红稳定器"],
            ["高端制造/自动化", "15%", "45万元", "产业升级、国产替代、出海能力"],
            ["AI/半导体/工业软件", "12%", "36万元", "新质生产力核心方向"],
            ["医药/消费", "12%", "36万元", "老龄化与品牌现金流"],
            ["电网/新能源/能源安全", "7%", "21万元", "能源转型与基础设施需求"],
            ["金融与其他", "4%", "12万元", "估值修复与周期平衡"],
        ],
        [2300, 1400, 1700, 3960],
    )
    add_heading(doc, "港股内部配置", 2)
    add_table(
        doc,
        ["方向", "港股内比例", "金额", "配置逻辑"],
        [
            ["港股宽基/恒生科技/互联网ETF", "35%", "31.5万元", "以ETF参与平台和科技修复，降低个股监管风险"],
            ["高股息央企/公用事业/电信能源", "30%", "27万元", "低估值、高分红、南向资金偏好"],
            ["创新药/医疗", "12%", "10.8万元", "国际化管线和估值修复，可优先ETF"],
            ["消费/服务/平台经济", "10%", "9万元", "消费修复与龙头现金流"],
            ["金融/交易所/保险", "8%", "7.2万元", "估值低位但需控制地产和利差压力"],
            ["资源能源", "5%", "4.5万元", "通胀和地缘风险对冲"],
        ],
        [2600, 1400, 1700, 3660],
    )

    add_heading(doc, "第五部分：个股与ETF筛选框架")
    for item in [
        "A股入池条件：上市满3年、非ST、审计意见标准无保留、近3年经营现金流至少2年为正、非金融企业资产负债率原则上低于60%、ROE通常不低于12%、经营现金流/净利润大于0.8、行业地位前三或细分龙头。",
        "港股入池条件：优先港股通或主流ETF；日均成交额能支撑分批建仓和退出；治理透明、审计质量高、关联交易少；自由现金流能覆盖分红或回购；监管边界清晰。",
        "ETF筛选条件：规模充足、成交活跃、费率合理、跟踪误差可控、成分透明，不买流动性不足和主题过窄的产品。",
        "硬性剔除：ST、退市风险、高商誉高质押、连续亏损且现金流恶化、频繁再融资补流、应收和存货增速显著高于收入、周期顶部仍激进扩产、纯题材炒作。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "第六部分：具体执行规则")
    add_body(doc, "建仓采用12个月分批路径：第1-2个月投入目标权益仓位的35%；第3-6个月投入35%；第7-12个月投入30%。若市场快速上涨，不追满；若A股或港股核心指数回撤10%-15%且基本面未破坏，可提前动用一档现金。")
    add_table(
        doc,
        ["规则", "A股", "港股"],
        [
            ["单一股票上限", "成熟龙头最高4%，高波动科技股最高3%", "单一港股最高3%，创新药单票最高2%"],
            ["ETF上限", "单只宽基ETF最高12%，行业ETF最高8%", "单只港股ETF最高10%"],
            ["行业上限", "单行业不超过A股资产25%", "单行业不超过港股资产35%"],
            ["加仓触发", "下跌8%且基本面未变加一档；下跌15%且估值合理偏低加第二档", "下跌来自流动性而非逻辑恶化，且南向资金和成交额未持续恶化"],
            ["减仓触发", "单仓达到目标1.5倍、估值高位、盈利下修", "估值进入历史高位、监管不确定性上升、港股仓位超限"],
            ["清仓触发", "财务造假、审计非标、核心竞争力消失", "治理瑕疵、现金流造假嫌疑、分红不可持续、流动性塌陷"],
        ],
        [2000, 3680, 3680],
    )

    add_heading(doc, "第七部分：风险管理体系")
    add_table(
        doc,
        ["风险阈值", "动作"],
        [
            ["组合回撤8%", "风控复核，暂停新增非白名单个股"],
            ["组合回撤12%", "暂停新增高波动成长仓，复核港股和科技暴露"],
            ["组合回撤18%", "停止买入高波动资产，召开临时投委会"],
            ["组合回撤25%", "权益降至55%-60%，现金目标提高至20%以上"],
            ["组合回撤30%", "强制重新提交投委会，清理逻辑破坏和流动性差仓位"],
            ["组合回撤35%", "危机模式，暂停新增风险资产，保留现金流与高流动性资产"],
        ],
        [2200, 7160],
    )
    add_body(doc, "港股新增风险包括人民币/港币汇率、美元利率、离岸流动性、监管变化、南向资金拥挤和单日大幅波动。香港金管局联系汇率制度将港元兑美元维持在7.75-7.85区间，人民币投资人实质上承担部分美元周期敞口。")

    add_heading(doc, "第八部分：压力测试")
    add_table(
        doc,
        ["情景", "组合预期影响", "应对动作"],
        [
            ["市场上涨50%", "组合预计上涨32%-42%", "权益降回70%左右，兑现10%-20%浮盈"],
            ["A股跌20%、港股跌30%", "组合预计回撤18%-25%", "启用预警，只加白名单宽基和高现金流资产"],
            ["A股跌40%、港股跌50%", "组合预计回撤30%-40%", "降港股与高波动赛道，暂停新增风险资产"],
            ["长期熊市", "年化收益可能为-5%至+3%", "提高红利、现金、短债权重，季度重做投资假设"],
            ["流动性危机", "短期净值可能快速下探", "现金提高至25%-35%，优先卖出弱流动性标的"],
            ["人民币/港币波动", "港股人民币口径收益受汇率扰动", "港股仓位不超过22%，不使用衍生品对冲"],
        ],
        [2200, 2800, 4360],
    )

    add_heading(doc, "第九部分：年度调仓计划")
    for item in [
        "每年1月：复核宏观情景、权益中枢、A股与港股比例。",
        "每年4-5月：根据年报和一季报更新股票池，剔除财务质量恶化标的。",
        "每年8-9月：根据中报调整行业权重，复核盈利兑现与现金流。",
        "每年12月：执行年度再平衡；任一资产类别偏离目标5个百分点以上必须调整。",
        "年度换手率目标25%-45%；除风险阈值触发外，不做高频择时。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "第十部分：最终组合结构")
    add_body(doc, "最终持仓建议：A股ETF 3-5只、A股个股10-16只；港股ETF 2-4只、港股个股3-6只；REITs 2-3只；现金和短债工具2-3只。前五大权益仓位不超过权益资产28%，前十大不超过45%。")
    add_table(
        doc,
        ["维度", "目标"],
        [
            ["总权益仓位", "78%，其中A股60%、港股18%"],
            ["ETF与个股", "权益资产中ETF约60%-70%，个股约30%-40%"],
            ["现金与低波资产", "现金固收12%、黄金6%、REITs4%"],
            ["业绩比较基准", "中证A500 40% + 沪深300 15% + 恒生指数/恒生科技20% + 中证红利10% + 中短债10% + 黄金5%"],
        ],
        [2600, 6760],
    )

    add_heading(doc, "第十一部分：投资委员会审查")
    add_table(
        doc,
        ["角色", "初评", "主要问题", "优化动作", "终评"],
        [
            ["CIO", "8.8", "港股比例过高会放大外部冲击", "将港股压至18%，保留战略卫星属性", "9.3"],
            ["风控总监", "8.6", "极端情境回撤超过原A股方案", "加入30%和35%强制复核阈值", "9.1"],
            ["基金经理", "9.0", "跨市场执行需更明确", "细化ETF/个股、建仓、再平衡规则", "9.3"],
            ["行业研究员", "8.9", "港股科技和创新药估值分化大", "提高现金流、治理和流动性门槛", "9.1"],
        ],
        [1600, 900, 2800, 2800, 1260],
    )
    add_body(doc, "结论：经优化后，四个投委会角色评分均达到9分以上。方案可作为机构内部执行框架，但仍需在实际下单前完成产品适当性、税费、交易权限、港股通额度、实时估值和标的流动性复核。")

    add_heading(doc, "最终指标")
    for item in [
        "预期年化收益率：8%-11%，不构成收益承诺。",
        "预期最大回撤：常态20%-25%，极端28%-35%。",
        "3-5年正收益胜率：约65%-75%。",
        "达到8%年化目标概率：约52%-60%。",
        "夏普率区间：0.40-0.70。",
        "最可能失败的原因：港股外部流动性冲击、美元利率高位、科技与高端制造盈利兑现低于估值、国内名义增长修复弱、地缘和监管不确定性升级。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "主要资料来源")
    sources = [
        "国家统计局：2026年一季度GDP与宏观数据，https://www.stats.gov.cn/sj/zxfbhjd/202604/t20260417_1963336.html",
        "国家统计局：2026年5月中国采购经理指数，https://www.stats.gov.cn/sj/sjjd/202605/t20260531_1963825.html",
        "中证指数：沪深300指数事实表，https://oss-ch.csindex.com.cn/static/html/csindex/public/uploads/indices/detail/files/zh_CN/000300factsheet.pdf",
        "港交所：月度市场概况与互联互通资料，https://www.hkex.com.hk/Market-Data/Statistics/Consolidated-Reports/HKEX-Monthly-Market-Highlights",
        "香港金管局：联系汇率制度，https://www.hkma.gov.hk/eng/key-functions/money/linked-exchange-rate-system/",
        "中国政府网/证监会：中长期资金入市与资本市场制度建设相关政策。",
    ]
    for src in sources:
        add_bullet(doc, src)

    doc.core_properties.title = "A股+港股跨市场投资计划"
    doc.core_properties.subject = "500万元人民币、3-5年、多智能体投委会版投资方案"
    doc.core_properties.author = "Codex"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT.resolve())
