from __future__ import annotations

from pathlib import Path

from docx import Document


path = Path("output/500万投资计划v7.0_A股港股具体股票投委会执行版_20260605.docx")
doc = Document(path)
texts = "\n".join(p.text for p in doc.paragraphs)
table_text = "\n".join(
    "\n".join("\t".join(cell.text for cell in row.cells) for row in table.rows)
    for table in doc.tables
)
all_text = texts + "\n" + table_text

required = [
    "一、核心参数",
    "二、完整配置表（总计500万）",
    "三、自动化风控系统",
    "四、执行决策树（时间优先 vs 价格优先）",
    "五、无衍生品对冲与再平衡",
    "六、止损自动化执行协议",
    "七、12周建仓计划",
    "八、行为保障体系",
    "九、预期收益与风险归因",
    "十、委员会反馈逐条回应",
    "十一、投委会审查",
    "美的集团 (000333.SZ)",
    "宁德时代 (300750.SZ)",
    "贵州茅台 (600519.SH)",
    "腾讯控股 (00700.HK)",
    "中国海洋石油 (00883.HK)",
    "信达生物 (01801.HK)",
    "不使用融资融券、不使用杠杆、不使用期货/期权/互换等衍生品",
    "不使用恒生指数PUT、股指期货、期权或其他衍生品",
    "A股权益\t300\t60%",
    "港股权益\t90\t18%",
]

print(f"exists={path.exists()} size={path.stat().st_size}")
print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")
for item in required:
    print(("OK   " if item in all_text else "MISS ") + item)

section = doc.sections[0]
print(
    "margins_twips="
    f"{section.top_margin.twips},"
    f"{section.right_margin.twips},"
    f"{section.bottom_margin.twips},"
    f"{section.left_margin.twips}"
)
for i, table in enumerate(doc.tables, 1):
    print(
        f"table_{i}=rows:{len(table.rows)},cols:{len(table.columns)},"
        f"header:{'|'.join(cell.text for cell in table.rows[0].cells)}"
    )
