from __future__ import annotations

from pathlib import Path

from docx import Document


path = Path(r"C:\Users\jjj64\WorkBuddy\20260318232610\500万A+H投资计划v8.0_最终执行版_20260605.docx")
doc = Document(path)

print(f"path={path}")
print(f"exists={path.exists()} size={path.stat().st_size}")
print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")

print("\nHEADINGS")
for idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    style = p.style.name if p.style else ""
    if style.startswith("Heading") or idx < 8:
        print(f"{idx:03d} [{style}] {text[:220]}")

print("\nTABLES")
for i, table in enumerate(doc.tables, 1):
    rows, cols = len(table.rows), len(table.columns)
    header = " | ".join(cell.text.replace("\n", " / ").strip() for cell in table.rows[0].cells) if rows else ""
    print(f"T{i}: rows={rows} cols={cols} header={header[:260]}")
    for r in range(min(rows, 5)):
        vals = " | ".join(cell.text.replace("\n", " / ").strip() for cell in table.rows[r].cells)
        print(f"  r{r}: {vals[:320]}")

all_text = "\n".join(p.text for p in doc.paragraphs) + "\n" + "\n".join(
    "\n".join("\t".join(cell.text for cell in row.cells) for row in t.rows)
    for t in doc.tables
)

checks = [
    "A股",
    "港股",
    "固收",
    "现金",
    "黄金",
    "REIT",
    "腾讯",
    "中海油",
    "美的",
    "宁德",
    "招商银行",
    "红利",
    "PUT",
    "期权",
    "AH",
    "南向",
    "再平衡",
    "加仓",
    "止损",
    "最大回撤",
    "收益",
    "投委会",
]
print("\nCHECKS")
for c in checks:
    print(f"{c}={all_text.count(c)}")
