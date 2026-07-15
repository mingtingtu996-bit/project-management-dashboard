from __future__ import annotations

from pathlib import Path

from docx import Document


path = Path(r"C:\Users\jjj64\WorkBuddy\20260318232610\500万A+H投资计划v8.0_最终执行版_20260605.docx")
doc = Document(path)

for idx in [3, 4, 5, 6, 21, 23, 24, 27, 28, 29, 31]:
    table = doc.tables[idx - 1]
    print(f"\n===== TABLE {idx} rows={len(table.rows)} cols={len(table.columns)} =====")
    for r, row in enumerate(table.rows):
        vals = [cell.text.replace("\n", " / ").strip() for cell in row.cells]
        print(f"r{r}: " + " | ".join(vals))

print("\n===== PUT PARAGRAPHS =====")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if "PUT" in text or "期权" in text or "对冲" in text:
        print(f"{i}: {text}")
