from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document


path = Path("output/A股港股跨市场投资计划_多智能体投委会版_20260605.docx")
doc = Document(path)
texts = "\n".join(p.text for p in doc.paragraphs)

required = [
    "多智能体工作流与审查结论",
    "第一部分：投资哲学与核心逻辑",
    "第二部分：未来3-5年中国投资主线",
    "第三部分：资产配置方案",
    "第四部分：行业配置方案",
    "第五部分：个股与ETF筛选框架",
    "第六部分：具体执行规则",
    "第七部分：风险管理体系",
    "第八部分：压力测试",
    "第九部分：年度调仓计划",
    "第十部分：最终组合结构",
    "第十一部分：投资委员会审查",
    "预期年化收益率：8%-11%",
    "预期最大回撤：常态20%-25%，极端28%-35%",
]

print(f"exists={path.exists()} size={path.stat().st_size}")
print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")
for item in required:
    print(("OK   " if item in texts else "MISS ") + item)

section = doc.sections[0]
print(
    "margins_twips="
    f"{section.top_margin.twips},"
    f"{section.right_margin.twips},"
    f"{section.bottom_margin.twips},"
    f"{section.left_margin.twips}"
)

with ZipFile(path) as package:
    xml = package.read("word/document.xml").decode("utf-8")
    footers = [
        package.read(name).decode("utf-8")
        for name in package.namelist()
        if name.startswith("word/footer") and name.endswith(".xml")
    ]

print(f"tblW_9360_count={xml.count('w:w=\"9360\"')}")
print(f"tblInd_count={xml.count('w:tblInd')}")
print(f"tcMar_count={xml.count('w:tcMar')}")
print(f"footer_count={len(footers)}")

for index, table in enumerate(doc.tables, start=1):
    print(
        f"table_{index}=rows:{len(table.rows)},cols:{len(table.columns)},"
        f"header:{'|'.join(cell.text for cell in table.rows[0].cells)}"
    )
