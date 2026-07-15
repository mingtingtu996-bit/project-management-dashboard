from __future__ import annotations

from pathlib import Path

from docx import Document


path = Path(r"C:\Users\jjj64\Desktop\500万投资计划v6.0_投委会通过版_20260605.docx")
doc = Document(path)

print(f"path={path}")
print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")
print("\nHEADINGS_AND_NONEMPTY_PARAGRAPHS")
for idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    style = p.style.name if p.style is not None else ""
    if style.startswith("Heading") or idx < 20 or text.startswith(("第", "一、", "二、", "三、", "四、", "五、", "六、", "七、", "八、", "九、", "十、")):
        print(f"{idx:03d} [{style}] {text}")

print("\nTABLES")
for i, table in enumerate(doc.tables, start=1):
    rows = len(table.rows)
    cols = len(table.columns)
    header = " | ".join(cell.text.replace("\n", " / ").strip() for cell in table.rows[0].cells) if rows else ""
    print(f"TABLE {i}: rows={rows} cols={cols} header={header}")
    for r in range(min(rows, 4)):
        vals = " | ".join(cell.text.replace("\n", " / ").strip() for cell in table.rows[r].cells)
        print(f"  r{r}: {vals}")
