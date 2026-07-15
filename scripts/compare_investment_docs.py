from __future__ import annotations

from pathlib import Path

from docx import Document


DOCS = {
    "user": Path(r"C:\Users\jjj64\WorkBuddy\20260318232610\500万A+H投资计划v7.0_投委会通过版_20260605.docx"),
    "codex": Path(r"C:\Users\jjj64\WorkBuddy\20260318232610\output\500万投资计划v7.0_A股港股具体股票投委会执行版_20260605.docx"),
}


def summarize(name: str, path: Path) -> None:
    doc = Document(path)
    print(f"===== {name} =====")
    print(f"path={path}")
    print(f"exists={path.exists()} size={path.stat().st_size}")
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")
    print("HEADINGS")
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        style = p.style.name if p.style is not None else ""
        if style.startswith("Heading") or idx < 8:
            print(f"{idx:03d} [{style}] {text[:180]}")
    print("TABLES")
    for i, table in enumerate(doc.tables, start=1):
        rows = len(table.rows)
        cols = len(table.columns)
        header = " | ".join(cell.text.replace("\n", " / ").strip() for cell in table.rows[0].cells) if rows else ""
        print(f"T{i}: rows={rows} cols={cols} header={header[:220]}")
        for r in range(min(rows, 3)):
            vals = " | ".join(cell.text.replace("\n", " / ").strip() for cell in table.rows[r].cells)
            print(f"  r{r}: {vals[:260]}")
    all_text = "\n".join(p.text for p in doc.paragraphs) + "\n" + "\n".join(
        "\n".join("\t".join(cell.text for cell in row.cells) for row in t.rows)
        for t in doc.tables
    )
    checks = [
        "PUT",
        "期权",
        "融资",
        "杠杆",
        "美的集团",
        "宁德时代",
        "贵州茅台",
        "腾讯控股",
        "中海油",
        "中国海洋石油",
        "招商银行",
        "北方华创",
        "中际旭创",
        "条件单",
        "12周",
        "最大回撤",
        "投委会",
    ]
    print("CHECKS")
    for c in checks:
        print(f"{c}={all_text.count(c)}")
    print()


for name, path in DOCS.items():
    summarize(name, path)
